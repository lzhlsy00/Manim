"""
File processing service for extracting text content from various file formats.
Supports PDF, Word documents, and images with OCR.
"""

import os
import io
import re
import tempfile
import logging
from typing import Optional, Dict, Any
from pathlib import Path
import mimetypes

# PDF processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
        USE_PDFPLUMBER = True
    except ImportError:
        PDF_AVAILABLE = False
        USE_PDFPLUMBER = False

# Word document processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Image OCR processing
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)

# File size limits (in bytes)
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
MAX_TEXT_LENGTH = 100000  # Maximum extracted text length

# Supported file types
SUPPORTED_FILE_TYPES = {
    'application/pdf': 'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/msword': 'doc',
    'image/jpeg': 'image',
    'image/png': 'image',
    'image/gif': 'image',
    'image/bmp': 'image',
    'image/tiff': 'image',
    'text/plain': 'text'
}


class FileProcessor:
    """File processing service for extracting text content."""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
        logger.info(f"File processor initialized with temp dir: {self.temp_dir}")
    
    def is_supported_file_type(self, content_type: str) -> bool:
        """Check if the file type is supported."""
        return content_type in SUPPORTED_FILE_TYPES
    
    def validate_file_size(self, file_size: int) -> bool:
        """Validate file size is within limits."""
        return file_size <= MAX_FILE_SIZE
    
    async def extract_text_from_file(self, file_content: bytes, filename: str, content_type: str) -> Optional[str]:
        """
        Extract text content from uploaded file based on its type.
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename
            content_type: MIME type of the file
            
        Returns:
            Extracted text content or None if extraction fails
        """
        try:
            file_type = SUPPORTED_FILE_TYPES.get(content_type)
            if not file_type:
                logger.warning(f"Unsupported file type: {content_type}")
                return None
            
            # Create temporary file
            temp_file_path = os.path.join(self.temp_dir, filename)
            with open(temp_file_path, 'wb') as temp_file:
                temp_file.write(file_content)
            
            text_content = None
            
            if file_type == 'pdf':
                text_content = await self._extract_from_pdf(temp_file_path)
            elif file_type in ['docx', 'doc']:
                text_content = await self._extract_from_word(temp_file_path)
            elif file_type == 'image':
                text_content = await self._extract_from_image(temp_file_path)
            elif file_type == 'text':
                text_content = await self._extract_from_text(temp_file_path)
            
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
            
            # Limit text length
            if text_content and len(text_content) > MAX_TEXT_LENGTH:
                text_content = text_content[:MAX_TEXT_LENGTH] + "\n\n[Text truncated due to length limit]"
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error extracting text from file {filename}: {str(e)}")
            return None
    
    async def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            logger.error("PDF processing not available. Install PyPDF2 or pdfplumber.")
            return None
        
        try:
            if USE_PDFPLUMBER:
                import pdfplumber
                text_content = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)
                return "\n\n".join(text_content)
            else:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_content = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)
                    return "\n\n".join(text_content)
                    
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            return None
    
    async def _extract_from_word(self, file_path: str) -> Optional[str]:
        """Extract text from Word document."""
        if not DOCX_AVAILABLE:
            logger.error("Word document processing not available. Install python-docx.")
            return None
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing Word document: {str(e)}")
            return None
    
    async def _extract_from_image(self, file_path: str) -> Optional[str]:
        """Extract text from image using OCR with enhanced math recognition."""
        # 运行时检测OCR可用性，避免导入时检测的问题
        try:
            import pytesseract
            from PIL import Image
        except ImportError as e:
            logger.error(f"OCR processing not available: {e}. Install pytesseract and pillow.")
            return None
        
        try:
            # Open image with PIL
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Try multiple OCR approaches for better math formula recognition
            text_results = []
            
            # Approach 1: Standard OCR with multiple languages
            try:
                text1 = pytesseract.image_to_string(image, lang='eng+chi_sim')
                if text1.strip():
                    text_results.append(("standard", text1.strip()))
            except Exception:
                pass
            
            # Approach 2: OCR optimized for math symbols
            try:
                # Use different OCR engine mode for math
                custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ+-=()[]{}∫∑∂∆πλμσθαβγδεζηθικλμνξοπρστυφχψω∞≤≥≠±×÷√∈∉⊂⊃⊆⊇∪∩∧∨¬→←↔↑↓'
                text2 = pytesseract.image_to_string(image, config=custom_config, lang='eng')
                if text2.strip():
                    text_results.append(("math_optimized", text2.strip()))
            except Exception:
                pass
            
            # Approach 3: Enhanced preprocessing and OCR
            try:
                # Resize image for better OCR accuracy
                width, height = image.size
                if width < 300 or height < 300:
                    new_size = (max(300, width * 2), max(300, height * 2))
                    resized_image = image.resize(new_size, Image.Resampling.LANCZOS)
                    text3 = pytesseract.image_to_string(resized_image, lang='eng+chi_sim')
                    if text3.strip():
                        text_results.append(("resized", text3.strip()))
            except Exception:
                pass
            
            # Select best result or combine them
            if text_results:
                # Log all attempts for debugging
                logger.info("=" * 60)
                logger.info("👁️ 【OCR处理详情】")
                logger.info(f"   📊 尝试方法数: {len(text_results)}")
                for method, text in text_results:
                    logger.info(f"   📝 {method}: {repr(text[:50])}...")
                
                # Use the longest result as it's likely more complete
                best_result = max(text_results, key=lambda x: len(x[1]))[1]
                logger.info(f"   ✨ 最佳结果: {repr(best_result[:50])}...")
                
                # Apply math symbol corrections
                corrected_text = self._correct_math_symbols(best_result)
                logger.info(f"   🔧 修正后结果: {repr(corrected_text[:100])}")
                logger.info("=" * 60)
                
                return corrected_text
            else:
                logger.warning("No text found in image with any OCR method")
                return None
            
        except Exception as e:
            logger.error(f"Error processing image with OCR: {str(e)}")
            return None
    
    def _correct_math_symbols(self, text: str) -> str:
        """Apply common math symbol corrections to OCR output."""
        # Common OCR misrecognitions for math symbols
        corrections = {
            # Integral symbols
            r'\[ve\s*\n?\s*x': '∫ x',  # Specific fix for the current issue
            r'\[': '∫',  # Left bracket often misrecognized as integral
            r'\\int': '∫',
            r'J\s*x': '∫ x',
            r'\]\s*x': '∫ x',
            
            # Common math symbols
            r'＋': '+',
            r'－': '-', 
            r'×': '×',
            r'÷': '÷',
            r'＝': '=',
            r'\bve\b': '',  # Remove common OCR noise
            r'\s+': ' ',  # Normalize whitespace
            
            # Exponents
            r'\^n': '^n',
            r'\bn\b': 'n',
            
            # Differential
            r'\bdx\b': 'dx',
            r'\bdy\b': 'dy',
            r'\bdt\b': 'dt',
        }
        
        corrected = text
        for pattern, replacement in corrections.items():
            corrected = re.sub(pattern, replacement, corrected, flags=re.IGNORECASE)
        
        # Clean up extra whitespace
        corrected = ' '.join(corrected.split())
        
        return corrected
    
    async def _extract_from_text(self, file_path: str) -> Optional[str]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encodings
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    return file.read()
            except Exception:
                try:
                    with open(file_path, 'r', encoding='latin-1') as file:
                        return file.read()
                except Exception as e:
                    logger.error(f"Error reading text file: {str(e)}")
                    return None
    
    def cleanup(self):
        """Clean up temporary directory."""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.info(f"Cleaned up temp directory: {self.temp_dir}")
        except Exception as e:
            logger.error(f"Error cleaning up temp directory: {str(e)}")
    
    def __del__(self):
        """Cleanup when object is destroyed."""
        self.cleanup()


# Global file processor instance
_file_processor = None

def get_file_processor() -> FileProcessor:
    """Get or create global file processor instance."""
    global _file_processor
    if _file_processor is None:
        _file_processor = FileProcessor()
    return _file_processor

def cleanup_file_processor():
    """Manually cleanup the global file processor."""
    global _file_processor
    if _file_processor is not None:
        _file_processor.cleanup()
        _file_processor = None